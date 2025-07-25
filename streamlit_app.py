import streamlit as st
import whisper
from docx import Document
import tempfile
import os

st.set_page_config(page_title="Transcriptor de Audio", layout="centered")

st.title("📝 Transcriptor de Audio con Whisper")
st.write("Sube un archivo de audio y genera su transcripción en texto y Word. Formatos compatibles: mp3, wav, m4a, mp4, aac.")

uploaded_file = st.file_uploader("📂 Sube tu archivo de audio", type=["mp3", "wav", "m4a", "mp4", "aac"])

if uploaded_file is not None:
    suffix = os.path.splitext(uploaded_file.name)[1]
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    word_file = "transcripcion.docx"

    try:
        with st.spinner("🔄 Transcribiendo audio, espera un momento..."):
            model = whisper.load_model("base")
            result = model.transcribe(tmp_path, language="es")

        st.success("✅ Transcripción completa")
        st.subheader("📄 Texto transcrito:")
        st.write(result["text"])

        doc = Document()
        doc.add_heading("Transcripción de Audio", 0)
        doc.add_paragraph(result["text"])
        doc.save(word_file)

        with open(word_file, "rb") as f:
            st.download_button("📥 Descargar Word", f, file_name=word_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        st.error(f"❌ Error al procesar el audio: {e}")

    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
        if os.path.exists(word_file):
            os.remove(word_file)

